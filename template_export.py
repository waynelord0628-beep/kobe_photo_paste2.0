from __future__ import annotations

import base64
import io
import re
import tempfile
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


EXPORT_MODES = {
    "two_landscape": "1頁2張橫式",
    "two_portrait": "1頁2張直式",
    "three_landscape": "1頁3張橫式",
    "four_portrait": "1頁4張直式",
    "six_portrait": "1頁6張直式",
}


def export_word(template_dir: Path, output_dir: Path, mode: str, title: str, photos: list[dict]) -> Path:
    if mode not in EXPORT_MODES:
        raise ValueError("不支援的模板版型")
    if not photos:
        raise ValueError("沒有照片可以匯出")

    title = sanitize_title(title or "Kobe強強照片黏貼")
    output_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="kobe_word_") as tmp:
        tmp_dir = Path(tmp)
        image_paths, names = materialize_images(photos, tmp_dir)
        if mode == "two_landscape":
            document = build_two_landscape(template_dir, title, image_paths, names)
        elif mode == "two_portrait":
            document = build_two_portrait(template_dir, title, image_paths, names)
        elif mode == "three_landscape":
            document = build_three_landscape(template_dir, title, image_paths, names)
        elif mode == "four_portrait":
            document = build_four_portrait(template_dir, title, image_paths, names)
        else:
            document = build_six_portrait(template_dir, title, image_paths, names)

        delete_first_paragraph_if_empty(document)
        delete_trailing_empty_paragraphs(document)
        target = unique_path(output_dir, title, ".docx")
        document.save(target)
        return target


def materialize_images(photos: list[dict], tmp_dir: Path) -> tuple[list[Path], list[str]]:
    paths: list[Path] = []
    names: list[str] = []
    for index, photo in enumerate(photos, start=1):
        url = str(photo.get("url") or "")
        if not url:
            continue
        image_bytes = decode_image_url(url)
        image = Image.open(io.BytesIO(image_bytes))
        if image.mode in ("RGBA", "P", "LA"):
            image = image.convert("RGBA")
        else:
            image = image.convert("RGB")
        path = tmp_dir / f"photo_{index:04d}.png"
        image.save(path, format="PNG")
        paths.append(path)
        names.append(photo_label(photo, index))
    return paths, names


def decode_image_url(url: str) -> bytes:
    if url.startswith("data:"):
        _, payload = url.split(",", 1)
        if ";base64" in url[: url.find(",")]:
            return base64.b64decode(payload)
        return unquote(payload).encode("utf-8")
    path = Path(url)
    if path.exists():
        return path.read_bytes()
    raise ValueError("照片資料不是可匯出的圖片格式")


def photo_label(photo: dict, index: int) -> str:
    description = str(photo.get("description") or "").strip()
    location = str(photo.get("location") or "").strip()
    return description or location or f"照片 {index}"


def build_two_landscape(template_dir: Path, title: str, image_paths: list[Path], names: list[str]) -> Document:
    document = open_template(template_dir / "word_template_portrait.docx")
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title)

    table = document.add_table(rows=len(image_paths) * 2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, (image_path, name) in enumerate(zip(image_paths, names)):
        cell_pic = table.cell(index * 2, 0)
        cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_pic = cell_pic.paragraphs[0]
        p_pic.add_run().add_picture(open_image_as_stream(image_path), width=Cm(15.5))
        p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_name = table.cell(index * 2 + 1, 0)
        fill_name_cell(cell_name, index + 1, f"說明：{name}", outer_width_cm=19.05)
        row_name = table.rows[index * 2 + 1]
        row_name.height = Cm(1.2)
        row_name.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for index in range(len(image_paths)):
        row_pic = table.rows[index * 2]
        row_pic.height = Cm(10.5)
        row_pic.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    return document


def build_two_portrait(template_dir: Path, title: str, image_paths: list[Path], names: list[str]) -> Document:
    document = open_template(template_dir / "word_template_portrait.docx")
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title)

    group_size = 2
    for group_start in range(0, len(image_paths), group_size):
        group_imgs = image_paths[group_start : group_start + group_size]
        group_names = names[group_start : group_start + group_size]
        table = document.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.rows[0].height = Cm(22.5)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[1].height = Cm(1.2)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        col_w = 9.525
        set_table_fixed_width(table, col_w * 2)
        for row in table.rows:
            for cell in row.cells:
                set_cell_width(cell, col_w)
        for offset, (image_path, name) in enumerate(zip(group_imgs, group_names)):
            cell_pic = table.cell(0, offset)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(image_path), width=Cm(8))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_name_cell(table.cell(1, offset), group_start + offset + 1, f"說明：{name}", outer_width_cm=col_w)
        if group_start + group_size < len(image_paths):
            document.add_section(WD_SECTION.NEW_PAGE)
    return document


def build_three_landscape(template_dir: Path, title: str, image_paths: list[Path], names: list[str]) -> Document:
    document = open_template(template_dir / "word_template_landscape.docx")
    section = document.sections[0]
    setup_header(document, section, title)

    group_size = 3
    for group_start in range(0, len(image_paths), group_size):
        group_imgs = image_paths[group_start : group_start + group_size]
        group_names = names[group_start : group_start + group_size]
        table = document.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.rows[0].height = Cm(15)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[1].height = Cm(1.2)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        col_w = 8.47
        set_table_fixed_width(table, col_w * 3)
        for row in table.rows:
            for cell in row.cells:
                set_cell_width(cell, col_w)
        for offset, (image_path, name) in enumerate(zip(group_imgs, group_names)):
            cell_pic = table.cell(0, offset)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(image_path), width=Cm(8))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_name_cell(table.cell(1, offset), group_start + offset + 1, f"說明：{name}", outer_width_cm=col_w)
        if group_start + group_size < len(image_paths):
            document.add_section(WD_SECTION.NEW_PAGE)
    return document


def build_four_portrait(template_dir: Path, title: str, image_paths: list[Path], names: list[str]) -> Document:
    document = open_template(template_dir / "word_template_portrait.docx")
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title)

    group_size = 4
    for group_start in range(0, len(image_paths), group_size):
        group_imgs = image_paths[group_start : group_start + group_size]
        group_names = names[group_start : group_start + group_size]
        table = document.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        col_w = 9.525
        set_table_fixed_width(table, col_w * 2)
        for row_index, row in enumerate(table.rows):
            row.height = Cm(10.2 if row_index in (0, 2) else 1.05)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if row_index in (0, 2) else WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                set_cell_width(cell, col_w)

        for offset, (image_path, name) in enumerate(zip(group_imgs, group_names)):
            row_base = 0 if offset < 2 else 2
            col = offset % 2
            cell_pic = table.cell(row_base, col)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(image_path), width=Cm(6.8))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_name_cell(table.cell(row_base + 1, col), group_start + offset + 1, f"說明：{name}", outer_width_cm=col_w)
        if group_start + group_size < len(image_paths):
            document.add_section(WD_SECTION.NEW_PAGE)
    return document


def build_six_portrait(template_dir: Path, title: str, image_paths: list[Path], names: list[str]) -> Document:
    document = open_template(template_dir / "word_template_portrait.docx")
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title)

    group_size = 6
    for group_start in range(0, len(image_paths), group_size):
        group_imgs = image_paths[group_start : group_start + group_size]
        group_names = names[group_start : group_start + group_size]
        table = document.add_table(rows=4, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        col_w = 6.35
        set_table_fixed_width(table, col_w * 3)
        for row_index, row in enumerate(table.rows):
            row.height = Cm(10.2 if row_index in (0, 2) else 1.0)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if row_index in (0, 2) else WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                set_cell_width(cell, col_w)

        for offset, (image_path, name) in enumerate(zip(group_imgs, group_names)):
            row_base = 0 if offset < 3 else 2
            col = offset % 3
            cell_pic = table.cell(row_base, col)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(image_path), width=Cm(4.45))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_name_cell(table.cell(row_base + 1, col), group_start + offset + 1, f"說明：{name}", outer_width_cm=col_w)
        if group_start + group_size < len(image_paths):
            document.add_section(WD_SECTION.NEW_PAGE)
    return document


def open_template(template_path: Path) -> Document:
    if not template_path.exists():
        raise FileNotFoundError(f"找不到模板檔：{template_path}")
    return Document(str(template_path))


def setup_header(document: Document, section, title: str) -> None:
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(0.3)
    header = section.header
    if header.paragraphs:
        paragraph = header.paragraphs[0]
        paragraph.clear()
    else:
        paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    run = paragraph.add_run(title)
    run.font.size = Pt(20)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    header.add_paragraph("")


def open_image_as_stream(image_path: Path) -> io.BytesIO:
    with Image.open(image_path) as image:
        if image.mode in ("RGBA", "P", "LA"):
            image = image.convert("RGBA")
        else:
            image = image.convert("RGB")
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer


def fill_name_cell(cell, number: int, desc_text: str, outer_width_cm: float = 9.0) -> None:
    from lxml import etree

    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)

    def make_paragraph(text: str, font_pt: int = 12) -> None:
        half_pts = str(font_pt * 2)
        p_el = etree.SubElement(tc, qn("w:p"))
        p_pr = etree.SubElement(p_el, qn("w:pPr"))
        spacing = etree.SubElement(p_pr, qn("w:spacing"))
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        r_el = etree.SubElement(p_el, qn("w:r"))
        r_pr = etree.SubElement(r_el, qn("w:rPr"))
        fonts = etree.SubElement(r_pr, qn("w:rFonts"))
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:eastAsia"), "標楷體")
        size = etree.SubElement(r_pr, qn("w:sz"))
        size.set(qn("w:val"), half_pts)
        size_cs = etree.SubElement(r_pr, qn("w:szCs"))
        size_cs.set(qn("w:val"), half_pts)
        t_el = etree.SubElement(r_el, qn("w:t"))
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_el.text = text

    desc_pt = best_font_size(desc_text, outer_width_cm)
    make_paragraph(f"編號 {number}")
    make_paragraph(desc_text, desc_pt)


def best_font_size(text: str, width_cm: float) -> int:
    pt_cm = 2.54 / 72
    usable = width_cm - 0.5
    for font_pt in (12, 11, 10, 9, 8):
        total = 0.0
        for char in text:
            total += font_pt * pt_cm if ord(char) > 0x2E7F else font_pt * pt_cm * 0.55
        if total <= usable:
            return font_pt
    return 8


def set_cell_width(cell, width_cm: float) -> None:
    from lxml import etree

    twips = cm_to_twips(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(old)
    tc_w = etree.SubElement(tc_pr, qn("w:tcW"))
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, total_width_cm: float) -> None:
    from lxml import etree

    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.SubElement(table._tbl, qn("w:tblPr"))
    for old in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old)
    tbl_w = etree.SubElement(tbl_pr, qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(cm_to_twips(total_width_cm)))
    tbl_w.set(qn("w:type"), "dxa")
    for old in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(old)
    layout = etree.SubElement(tbl_pr, qn("w:tblLayout"))
    layout.set(qn("w:type"), "fixed")


def cm_to_twips(value: float) -> int:
    return int(value / 2.54 * 1440)


def delete_first_paragraph_if_empty(document: Document) -> None:
    if document.paragraphs and not document.paragraphs[0].text.strip():
        element = document.paragraphs[0]._element
        element.getparent().remove(element)


def delete_trailing_empty_paragraphs(document: Document) -> None:
    body = document.element.body
    while True:
        children = list(body)
        if not children:
            return
        last = children[-1]
        if last.tag == qn("w:sectPr") or last.tag != qn("w:p"):
            return
        text = "".join(t.text or "" for t in last.iter(qn("w:t"))).strip()
        has_drawing = last.find(".//" + qn("w:drawing")) is not None
        if text or has_drawing:
            return
        body.remove(last)


def sanitize_title(title: str) -> str:
    safe = re.sub(r'[\\/*?:"<>|]', "_", title).strip()
    return safe or "Kobe強強照片黏貼"


def unique_path(folder: Path, base_name: str, suffix: str) -> Path:
    target = folder / f"{base_name}{suffix}"
    index = 1
    while target.exists():
        target = folder / f"{base_name}_{index}{suffix}"
        index += 1
    return target
